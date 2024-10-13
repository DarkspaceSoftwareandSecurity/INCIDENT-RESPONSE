import os
import platform
import subprocess
import socket
import cv2
from docx import Document
import openpyxl
import pyshark
import datetime
import matplotlib.pyplot as plt

# Author Information
AUTHOR = "Michael James Blenkinsop"
EMAIL = "mickyblenk@gmail.com"
PHONE = "+447710669684"
WEBSITE = "darkspacesoftwareandsecurity.com"

# Get the base directory dynamically from where the script is run
BASE_DIR = os.getcwd()
REPORT_DIR = os.path.join(BASE_DIR, "reports")
LOGS_DIR = os.path.join(BASE_DIR, "logs")
PCAP_DIR = os.path.join(BASE_DIR, "pcap")

# Ensure necessary folders are created dynamically
os.makedirs(REPORT_DIR, exist_ok=True)
os.makedirs(LOGS_DIR, exist_ok=True)
os.makedirs(PCAP_DIR, exist_ok=True)

# Display DarkSpace Software and Security Banner
def display_banner():
    print("\n" + "=" * 60)
    print("  DARKSPACE SOFTWARE & SECURITY")
    print("=" * 60)
    print("© 2024 DarkSpace Software & Security. All rights reserved.")
    print(f"Author: {AUTHOR}")
    print(f"Email: {EMAIL}")
    print(f"Phone: {PHONE}")
    print(f"Website: {WEBSITE}")
    print("=" * 60 + "\n")

# Install Dependencies Function
def install_dependencies():
    try:
        requirements = [
            "opencv-python",
            "python-docx",
            "openpyxl",
            "matplotlib",
            "pyshark"
        ]

        for package in requirements:
            subprocess.check_call([os.sys.executable, "-m", "pip", "install", package])
        
        print("Dependencies installed successfully.")
    except subprocess.CalledProcessError as e:
        print(f"Error installing dependencies: {e}")

# Network Monitoring and Blocking Function (IDS/EDR/XDR)
def monitor_network(mode):
    try:
        suspicious_ips = ['192.168.1.50']  # Example IPs for demonstration
        with socket.socket(socket.AF_INET, socket.SOCK_RAW, socket.IPPROTO_IP) as sniffer:
            sniffer.bind((socket.gethostname(), 0))
            print(f"Network Monitoring Started in {mode} mode. Press Ctrl+C to stop.")
            while True:
                packet = sniffer.recvfrom(65565)
                ip = packet[1][0]
                if ip in suspicious_ips:
                    os.system(f"route delete {ip}")
                    with open(os.path.join(LOGS_DIR, 'network_log.txt'), 'a') as log_file:
                        log_file.write(f"{datetime.datetime.now()} - Blocked suspicious IP ({mode}): {ip}\n")
                    print(f"Blocked suspicious IP ({mode}): {ip}")
    except KeyboardInterrupt:
        print(f"Network monitoring stopped for {mode} mode.")
    except Exception as e:
        print(f"Network monitoring error: {e}")

# Packet Capture Function (PCAP)
def start_packet_capture(duration=60):
    try:
        pcap_file_path = os.path.join(PCAP_DIR, f"capture_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.pcap")
        print(f"Starting packet capture for {duration} seconds. Saving to {pcap_file_path}.")
        capture = pyshark.LiveCapture(output_file=pcap_file_path)
        capture.sniff(timeout=duration)
        print(f"Packet capture saved to {pcap_file_path}.")
        return pcap_file_path
    except Exception as e:
        print(f"Error capturing packets: {e}")
        return None

# Capture Image from Attacker's Camera Function
def capture_attacker_image():
    try:
        cap = cv2.VideoCapture(0)
        if cap.isOpened():
            ret, frame = cap.read()
            if ret:
                image_path = os.path.join(REPORT_DIR, "attacker_image.jpg")
                cv2.imwrite(image_path, frame)
                cap.release()
                print(f"Image captured successfully and saved at {image_path}")
                return image_path
        else:
            print("Failed to access the camera.")
    except Exception as e:
        print(f"Error capturing attacker image: {e}")
    finally:
        cv2.destroyAllWindows()
    return None

# Generate Forensic Report Function
def generate_report(attack_data, image_path=None, pcap_file=None):
    try:
        # Create DOCX report
        doc = Document()
        doc.add_heading('Forensic Incident Report', 0)
        doc.add_paragraph(f"Attack Details: {attack_data}")
        if image_path:
            doc.add_paragraph("Image Evidence:")
            doc.add_picture(image_path, width=2000000)  # Adding an image with a fixed width
        if pcap_file:
            doc.add_paragraph(f"Packet Capture File: {pcap_file}")

        doc_path = os.path.join(REPORT_DIR, 'forensic_report.docx')
        doc.save(doc_path)

        # Create XLS report
        wb = openpyxl.Workbook()
        sheet = wb.active
        sheet.title = 'Attack Data'
        sheet['A1'] = 'Details'
        sheet['A2'] = attack_data
        xls_path = os.path.join(REPORT_DIR, 'forensic_data.xlsx')
        wb.save(xls_path)

        print(f"Reports generated successfully: {doc_path}, {xls_path}")
        return doc_path, xls_path
    except Exception as e:
        print(f"Error generating report: {e}")
        return None, None

# Graphical Display of Attack Patterns
def plot_attack_trace(data):
    try:
        plt.figure(figsize=(10, 6))
        plt.bar(data.keys(), data.values(), color='purple')
        plt.xlabel('IP Address')
        plt.ylabel('Attack Frequency')
        plt.title('Attack Pattern Analysis')
        plt.grid(True)
        plt.show()
    except Exception as e:
        print(f"Error plotting attack trace: {e}")

# Clear Screen Function
def clear_screen():
    if platform.system() == "Windows":
        os.system('cls')
    else:
        os.system('clear')

# User Voice Command Interaction with Menu
def user_voice_interaction():
    mode = "IDS"  # Default mode
    while True:
        print(f"\nGreetings, Commander. You are currently in {mode} mode.")
        print("\nType 'show options' at any time to see what actions you can take.")

        command = input("[COMMANDER] > ").strip().lower()
        
        if command == "show options":
            print("\nAvailable Actions, Commander:")
            print("1. Switch Modes (IDS/EDR/XDR)")
            print("2. Start Network Monitoring")
            print("3. Start Packet Capture")
            print("4. Capture Attacker Image")
            print("5. Generate Forensic Report")
            print("6. View Attack Patterns (Demo)")
            print("7. Exit")
        
        elif command == 'clear' or command == 'cls':
            clear_screen()
            print("\nWhat would you like to do next, Commander?")
            print("Available Actions:")
            print("1. Switch Modes (IDS/EDR/XDR)")
            print("2. Start Network Monitoring")
            print("3. Start Packet Capture")
            print("4. Capture Attacker Image")
            print("5. Generate Forensic Report")
            print("6. View Attack Patterns (Demo)")
            print("7. Exit")

        elif command == '1':
            print("\nWhich mode would you like to switch to, Commander? (IDS/EDR/XDR)")
            new_mode = input("[COMMANDER] > ").strip().upper()
            if new_mode in ["IDS", "EDR", "XDR"]:
                mode = new_mode
                print(f"Switched to {mode} mode, Commander.")
            else:
                print("Invalid mode, Commander. Please choose IDS, EDR, or XDR.")
        
        elif command == '2':
            print(f"\n[INFO] Starting Network Monitoring in {mode} mode, Commander.")
            monitor_network(mode)
        
        elif command == '3':
            print("\n[INFO] Starting Packet Capture, Commander.")
            duration = input("Enter packet capture duration in seconds (default 60): ")
            duration = int(duration) if duration.isdigit() else 60
            start_packet_capture(duration)
        
        elif command == '4':
            print("\n[INFO] Capturing Attacker Image, Commander.")
            capture_attacker_image()
        
        elif command == '5':
            print("\n[INFO] Generating Forensic Report, Commander.")
            attack_data = "Example attack detected..."
            image_path = os.path.join(REPORT_DIR, "attacker_image.jpg") if os.path.exists(os.path.join(REPORT_DIR, "attacker_image.jpg")) else None
            pcap_files = [f for f in os.listdir(PCAP_DIR) if f.endswith('.pcap')]
            pcap_file = os.path.join(PCAP_DIR, pcap_files[-1]) if pcap_files else None
            generate_report(attack_data, image_path, pcap_file)
        
        elif command == '6':
            print("\n[INFO] Displaying Attack Patterns, Commander.")
            sample_data = {'192.168.1.50': 5, '192.168.1.51': 3, '192.168.1.52': 7}
            plot_attack_trace(sample_data)
        
        elif command == '7':
            print("\n[INFO] Exiting Response AI, Commander. Farewell.")
            break
        
        else:
            print("\n[ERROR] Invalid command, Commander. Type 'show options' to see available actions.")

# Main Execution
if __name__ == "__main__":
    # Display Banner
    display_banner()

    # Step 1: Install dependencies
    install_dependencies()

    # Step 2: Start User Voice Interaction
    user_voice_interaction()
